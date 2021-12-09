import hashlib
import base64
import re
import datetime
import os
import time
import numbers
import json
import tempfile
import mimetypes
import numpy as np
import pandas as pd
import bokeh
from glob import glob
from bokeh.models import Title
from bokeh.plotting import figure 
from bokeh.embed import components
#from bokeh.models.tools import HoverTool
from bokeh.models import HoverTool, TapTool, CustomJS, LogTicker, NumeralTickFormatter #,PanTool
from bokeh.models import BasicTicker, ColorBar, ColumnDataSource, LinearColorMapper, PrintfTickFormatter
from bokeh.transform import transform
from bokeh.embed import json_item
import xml.etree.ElementTree as ET
from copy import deepcopy
from html.parser import HTMLParser

import urllib.request
from urllib.parse import urlparse
import docx
from docx import Document
from htmldocx import HtmlToDocx
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

from .rdkit import TCFingerPrint

from .utils import order_queryset_list_by_values_list

from django.db import IntegrityError
from django.core.exceptions import ObjectDoesNotExist
from django.conf import settings
from django.http import HttpResponse, JsonResponse, response
from django.shortcuts import render, redirect
from rest_framework import status
from django.contrib.auth import get_user_model
from django.contrib.auth import authenticate, login, logout as auth_logout
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from django.views.decorators.http import require_http_methods
from django.views.decorators.csrf import csrf_protect, ensure_csrf_cookie
from django.middleware.csrf import get_token
from django.utils import timezone
from django.db.models import F, Max
from django.db import transaction
from django.db import OperationalError
from django.http import Http404
from django.utils.html import escape

from rest_framework.decorators import api_view
from rest_framework import serializers, status

from rest_framework.views import APIView
from rest_framework.generics import GenericAPIView, RetrieveUpdateDestroyAPIView, ListAPIView, ListCreateAPIView
from rest_framework.mixins import UpdateModelMixin, CreateModelMixin, RetrieveModelMixin, ListModelMixin
from rest_framework.response import Response
from rest_framework.authentication import SessionAuthentication
from rest_framework.permissions import IsAuthenticated
from rest_framework.parsers import MultiPartParser
from rest_framework.exceptions import PermissionDenied, server_error

from .models import Projects as ProjectsModel, ProblemDescription, InitialRAxHypothesis
from .models import Nodes as NodesModel, NodeType
from .models import Resources as ResourcesModel
from .models import File, FileType
from .models import Compound, DataMatrix, DataMatrixFields, UnitType, Unit, CompoundCASRN, TCompound

from .serializer import ProjectSerializer, UserSerializer, NodeSerializer, FullNodeSerializer
from .serializer import StatusSerializer, ResourcesSerializer, ProblemDescriptionSerializer, InitialRAxHypothesisSerializer, SlugCompoundCASRNSSerializerNoValidation
from .serializer import CompoundSerializer, DataMatrixSerializer, UnitTypeSerializer, UnitSerializer, CompoundCASRNSerializer, SlugCompoundCASRNSSerializer
from .serializer import DataMatrixFieldsSerializer, DataMatrixFieldsReadSerializer, ChemblDataMatrixSerializer, CompoundDataMatrixSerializer, CompoundSerializer, TCompoundSerializer
from .serializer import CompoundImageImageSerializer, CompoundDataMatrixCountSerializer
from .permissions import IsProjectOwner


from django.utils.decorators import available_attrs
from functools import wraps

base54re = re.compile(r'^data:(.*);base64,(.*)$')

@method_decorator((csrf_protect,ensure_csrf_cookie), name='dispatch')
class ManageProject(RetrieveUpdateDestroyAPIView):
    """
    Add Project
    """
    queryset = ProjectsModel.objects.all()
    serializer_class = ProjectSerializer
    lookup_field = 'project'
    lookup_url_kwarg = 'project'
    authentication_classes = [SessionAuthentication]
    permission_classes = [IsAuthenticated,IsProjectOwner]



@method_decorator((csrf_protect,ensure_csrf_cookie), name='dispatch')
class ProjectStatus(ListAPIView):
    """
    Add Project
    """
    queryset = NodesModel.objects.values('node_seq','executed')
    serializer_class = StatusSerializer
    lookup_field = 'project'
    lookup_url_kwarg = 'project'
    authentication_classes = [SessionAuthentication]
    permission_classes = [IsAuthenticated,IsProjectOwner]

    def get_queryset(self):
        return self.queryset.filter(**{self.lookup_field:self.kwargs[self.lookup_url_kwarg]})

@method_decorator((csrf_protect,ensure_csrf_cookie), name='dispatch')
class ListProjects(ListAPIView):
    serializer_class = ProjectSerializer
    authentication_classes = [SessionAuthentication]
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        projects = ProjectsModel.objects.filter(owner=self.request.user.id)
        return projects
 

@method_decorator((csrf_protect,ensure_csrf_cookie), name='dispatch')
class ManageNodes(GenericAPIView,UpdateModelMixin,CreateModelMixin):
    serializer_class = NodeSerializer
    lookup_field = 'project'
    lookup_url_kwarg = 'project'
    authentication_classes = [SessionAuthentication]
    permission_classes = [IsAuthenticated,IsProjectOwner]
    
    def get_queryset(self):
        return NodesModel.objects.filter(project=self.kwargs['project'], node_seq=self.kwargs['node'])

    def get(self, request, project, node):
        not_found = False
        newdict = {}
        newdict['inputs'] = []
        resources = ResourcesModel.objects.filter(node=node)
        resources = ResourcesSerializer(resources, many=True).data
        newdict['resources'] = resources
        try:
            node_info = NodesModel.objects.annotate(name=F('node_seq__name'),description=F('node_seq__description'), history_node_list=F('node_seq__history_node_list')).get(project=project,node_seq=node)
            node_info_history_node_list = node_info.history_node_list
            newdict.update(FullNodeSerializer(node_info, many=False).data)
        except NodesModel.DoesNotExist:
            newdict['Reason'] = 'Node not found.'
            node_info = list(NodeType.objects.filter(id=node).values('name','description','history_node_list'))
            if len(node_info) == 0:
                newdict['Reason'] = 'Node type not found.'
                return Response(newdict, status.HTTP_404_NOT_FOUND)
            node_info = node_info[0]
            node_info['project'] = int(project)
            node_info['node_seq'] = int(node)
            node_info_history_node_list = node_info['history_node_list']
            newdict.update(node_info)

        
        qhistory = NodesModel.objects.filter(project=project)
        qhistory = qhistory.annotate(content=F('outputs'),comment=F('outputs_comments'), name=F('node_seq__name')).values('name','content','comment','inputs_comments','node_seq')

        if node_info_history_node_list is None:
            qhistory = qhistory.filter(node_seq__lt=node).order_by('node_seq')
            history = list(qhistory)
        elif node_info_history_node_list == '':
            history = []
        else:
            history_node_list_as_list = list(map(int,node_info_history_node_list.split(',')))
            qhistory = qhistory.filter(node_seq__in=history_node_list_as_list)
            history = order_queryset_list_by_values_list(qhistory,'node_seq', history_node_list_as_list)
        
        for i in range(0,len(history)):
            histi = dict(history[i])
            content = histi['content']
            #if content != "":
            #    content = "<b>Output:</b><br>"+histi['content']
            # if history[i]['node_seq'] == 1:
            #     try:
            #         qproblem = ProblemDescription.objects.get(project=project)
            #         #histi['content'] = "<b>Problem description:</b><br>"+qproblem.description+content
            #         histi['content'] = qproblem.description
            #     except ProblemDescription.DoesNotExist:
            #         histi['content'] = content
            #else:
            #    if histi['inputs_comments'] != "":
            #        histi['content'] = "<b>Input:</b><br>"+histi['inputs_comments']+content
            #    else:
            #        histi['content'] = content
            histi.pop('inputs_comments', None)
            histi.pop('node_seq', None)
            history[i] = histi

        newdict['inputs'] = history
        newdict['CSRF_TOKEN'] = get_token(request)
        if not_found: 
            return Response(newdict, status.HTTP_404_NOT_FOUND)
        else:
            return Response(newdict, status.HTTP_200_OK)

    def post(self, request, project, node):
        try:
            with transaction.atomic():
                response = self.partial_update(request)
                if response.status_code == status.HTTP_200_OK:
                    self.get_queryset().update(executed=True)
        except Http404 as e:
            data = dict(request.data)
            for key in data:
                data[key] = data[key][0]
            data['project'] = int(project)
            data['node_seq'] = int(node)
            data['executed'] = True 
            serializer = self.get_serializer(data=data)
            serializer.is_valid(raise_exception=True)
            self.perform_create(serializer)
            headers = self.get_success_headers(serializer.data)
            #return Response(serializer.data, status=status.HTTP_201_CREATED, headers=headers)

        return JsonResponse({'Ok':'ok'}, status=status.HTTP_200_OK)

@method_decorator((csrf_protect,ensure_csrf_cookie), name='dispatch')
class ProblemDescriptionView(GenericAPIView,CreateModelMixin,RetrieveModelMixin,UpdateModelMixin):
    serializer_class = ProblemDescriptionSerializer
    lookup_field = 'project'
    lookup_url_kwarg = 'project'
    authentication_classes = [SessionAuthentication]
    permission_classes = [IsAuthenticated,IsProjectOwner]
    queryset = ProblemDescription.objects.all()
    
    def get(self, request, project):
        try:
            response = self.retrieve(request)
        except Http404:
            return Response({'description':'Not Found: '+request.path,'CSRF_TOKEN':get_token(request)}, status=status.HTTP_404_NOT_FOUND)
        response.data['CSRF_TOKEN'] = get_token(request)
        return response

    def post(self, request, project):
        data = dict(request.data)
        for key in data:
            data[key] = data[key][0]
        data['project'] = project
        
        # Perform the lookup filtering.
        lookup_url_kwarg = self.lookup_url_kwarg or self.lookup_field

        assert lookup_url_kwarg in self.kwargs, (
            'Expected view %s to be called with a URL keyword argument '
            'named "%s". Fix your URL conf, or set the `.lookup_field` '
            'attribute on the view correctly.' %
            (self.__class__.__name__, lookup_url_kwarg)
        )
        filter_kwargs = {self.lookup_field: self.kwargs[lookup_url_kwarg]}
        try:
            instance = self.get_queryset().get(**filter_kwargs)
            serializer = self.get_serializer(instance, data=data, partial=True)
        except ObjectDoesNotExist as e:
            serializer = self.get_serializer(data=data)
        serializer.is_valid(raise_exception=True)
        del data
        data = dict(serializer.validated_data)
        project_obj = data.pop('project')
        success = False
        while not success:
            try:
                obj, created = self.get_queryset().update_or_create(project=project_obj, defaults=data)
                success = True
            except OperationalError as e:
                if str(e) == "database is locked":
                    time.sleep(5)
                else:
                    print(e)
                    raise e

        if created:
            NodesModel.objects.filter(project=self.kwargs['project'], node_seq=1).update(executed=True)
        
        return Response({'Ok':'ok'}, status=status.HTTP_200_OK)
    
@method_decorator((csrf_protect,ensure_csrf_cookie), name='dispatch')
class InitialRAxHypothesisView(GenericAPIView,CreateModelMixin,RetrieveModelMixin,UpdateModelMixin):
    serializer_class = InitialRAxHypothesisSerializer
    lookup_field = 'project'
    lookup_url_kwarg = 'project'
    authentication_classes = [SessionAuthentication]
    permission_classes = [IsAuthenticated,IsProjectOwner]
    queryset = InitialRAxHypothesis.objects.all()
    
    def get(self, request, project):
        try:
            response = self.retrieve(request)
        except Http404:
            return Response({'description':'Not Found: '+request.path,'CSRF_TOKEN':get_token(request)}, status=status.HTTP_404_NOT_FOUND)
        response.data['CSRF_TOKEN'] = get_token(request)
        return response

    def post(self, request, project):
        data = dict(request.data)
        for key in data:
            data[key] = data[key][0]
        data['project'] = project
        
        # Perform the lookup filtering.
        lookup_url_kwarg = self.lookup_url_kwarg or self.lookup_field

        assert lookup_url_kwarg in self.kwargs, (
            'Expected view %s to be called with a URL keyword argument '
            'named "%s". Fix your URL conf, or set the `.lookup_field` '
            'attribute on the view correctly.' %
            (self.__class__.__name__, lookup_url_kwarg)
        )
        filter_kwargs = {self.lookup_field: self.kwargs[lookup_url_kwarg]}
        try:
            instance = self.get_queryset().get(**filter_kwargs)
            serializer = self.get_serializer(instance, data=data, partial=True)
        except ObjectDoesNotExist as e:
            serializer = self.get_serializer(data=data)
        serializer.is_valid(raise_exception=True)
        del data
        data = dict(serializer.validated_data)
        project_obj = data.pop('project')
        success = False
        while not success:
            try:
                obj, created = self.get_queryset().update_or_create(project=project_obj, defaults=data)
                success = True
            except OperationalError as e:
                if str(e) == "database is locked":
                    time.sleep(5)
                else:
                    print(e)
                    raise e

        if created:
            NodesModel.objects.filter(project=self.kwargs['project'], node_seq=1).update(executed=True)
        return Response({'Ok':'ok'}, status=status.HTTP_200_OK)
    
@method_decorator((csrf_protect,ensure_csrf_cookie), name='dispatch')
class User(APIView):
    # If the user is already logged in, it responds a JSON with user data and the CSRF token.
    # Otherwise, only responds a JSON with the CSRF token for the POST method.
    def get(self,request,logout):
        respdata = {}
        if request.user.is_authenticated:
            respdata = UserSerializer(request.user, many=False).data
        respdata['CSRF_TOKEN'] = get_token(request) 
        return JsonResponse(respdata, status=status.HTTP_200_OK)

    def post(self,request,logout):
        if logout == "logout":
            if request.user.is_authenticated:
                auth_logout(request)
                request.session.pop('rememberme', default=None)
                return JsonResponse({'Ok':'ok'}, status=status.HTTP_200_OK)
        else:
            
            if 'username' in request.POST or 'password' in request.POST:
                username = request.POST.get('username')
                password = request.POST.get('password')

                user = authenticate(username=username, password=password)
                if user is not None:
                    login(request, user)
            
            if request.user.is_authenticated:
                if 'rememberme' in request.POST:
                    if int(request.POST.get('rememberme')):
                        request.session['rememberme'] = True
                        request.session.set_expiry(settings.SESSION_COOKIE_AGE)
                    else:
                        request.session['rememberme'] = False
                        request.session.set_expiry(0)
                else:
                    request.session.set_expiry(0)
                return Response(UserSerializer(request.user, many=False).data, status.HTTP_200_OK)

        raise PermissionDenied


class Resources(ListAPIView):
    queryset = ResourcesModel.objects.all()
    serializer_class = ResourcesSerializer
    lookup_field = 'node'
    lookup_url_kwarg = 'node'

    def get_queryset(self):
        return self.queryset.filter(**{self.lookup_field:self.kwargs[self.lookup_url_kwarg]})

@method_decorator((csrf_protect,ensure_csrf_cookie), name='dispatch')
class FileUploadView(APIView):
    parser_classes = [MultiPartParser]
    authentication_classes = [SessionAuthentication]
    permission_classes = [IsAuthenticated,IsProjectOwner]

    # save_uploadedfile function from https://github.com/GPCRmd/GPCRmd
    # Copyright (c) 2017 Ismael Rodríguez-Espigares et al., Jana Selent,
    # UPF and IMIM
    # Apache License 2.0 (http://www.apache.org/licenses/LICENSE-2.0)
    def _save_uploadedfile(self,filepath,uploadedfile):
        with open(filepath,'wb') as f:
            if uploadedfile.multiple_chunks:
                for chunk in uploadedfile.chunks():
                    f.write(chunk)
            else:
                f.write(uploadedfile.read())
            f.close()

    def get(self, request, project, node, part):
        respdata = {'msg': 'OK'}
        respdata['CSRF_TOKEN'] = get_token(request) 
        return JsonResponse(respdata, status=status.HTTP_200_OK)

    def post(self, request, project, node, part, format=None):
        filekey = 'file'
        if filekey not in request.FILES:
            return Response({"detail","%s file field is missing." % (filekey)},status=status.HTTP_400_BAD_REQUEST)
        file_list = request.FILES.getlist(filekey)
        if len(file_list) > 1:
            return Response({"detail","Only one file in %s is acceptable." % (filekey)},status=status.HTTP_400_BAD_REQUEST)
        elif len(file_list) == 0:
            return Response({"detail","%s has no files." % (filekey)},status=status.HTTP_400_BAD_REQUEST)

        uploadedfile = file_list[0]
        filename = uploadedfile.name
        rootname,fileext = os.path.splitext(uploadedfile.name)
        timestamp = timezone.now()
        date = timestamp.date()
        micro = timestamp.microsecond

        # create a MD5 hash of filename+(micosecond current time fraction) and encode it to base32 removing
        # '=' padding for file names compatible with Windows.
        md5 = base64.b32encode(hashlib.md5((filename+str(micro)).encode()).digest()).decode().strip('=')
        newfilename = '_'.join((date.strftime('%d%m%Y'),str(project),str(node),str(part),str(md5)))
        folderpath = os.path.join('projects',str(project),'node_'+str(node))
        
        fullfolderpath = os.path.join(settings.MEDIA_ROOT_UPLOADS,folderpath)
        os.makedirs(fullfolderpath,mode=0o770,exist_ok=True)

        
        # TODO optimize with filter and anotate
        node_obj = NodesModel.objects.get(project=project,node_seq=node)

        # TODO guess file type by extension
        # TODO check if file type is accepted
        file_type = 1
        
        file_type_obj = FileType.objects.get(pk=file_type)
        colindex = 0
        while 1:
            try:
                filepath = os.path.join(fullfolderpath,newfilename+fileext)
                self._save_uploadedfile(filepath,uploadedfile)
                url = os.path.join(settings.MEDIA_ROOT_UPLOADS,folderpath,newfilename+fileext)
                new_obj, created = File.objects.get_or_create(user_filename=filename,node=node_obj,part=part,
                    defaults= {'filename':newfilename+fileext,'file_type':file_type_obj,
                    'filepath':filepath,'url':url})
            except IntegrityError as e:
                obj = File.objects.get(user_filename=filename,node=node_obj,part=part)
                if obj.filename == newfilename+fileext:
                    newfilename += '_'+str(colindex)
                    colindex += 1
            except Exception:
                raise
            else:
                break
        if not created:
            return Response(data={'detail':'File "%s" already exists.' % (filename),'URL':new_obj.url},status=400)
        return Response(data={'msg':'OK','URL':new_obj.url},status=status.HTTP_200_OK)

@method_decorator((csrf_protect,ensure_csrf_cookie), name='dispatch')
class CSVFileToHTML(APIView):
    parser_classes = [MultiPartParser]
    authentication_classes = [SessionAuthentication]
    permission_classes = [IsAuthenticated,IsProjectOwner]

    # save_uploadedfile function from https://github.com/GPCRmd/GPCRmd
    # Copyright (c) 2017 Ismael Rodríguez-Espigares et al., Jana Selent,
    # UPF and IMIM
    # Apache License 2.0 (http://www.apache.org/licenses/LICENSE-2.0)
    def _save_uploadedfile(self,filepath,uploadedfile):
        with open(filepath,'wb') as f:
            if uploadedfile.multiple_chunks:
                for chunk in uploadedfile.chunks():
                    f.write(chunk)
            else:
                f.write(uploadedfile.read())
            f.close()

    def get(self, request, project, node, part):
        respdata = {'msg': 'OK'}
        respdata['CSRF_TOKEN'] = get_token(request) 
        return JsonResponse(respdata, status=status.HTTP_200_OK)

    def post(self, request, project, node, part, format=None):
        filekey = 'file'
        if filekey not in request.FILES:
            return Response({"detail","%s file field is missing." % (filekey)},status=status.HTTP_400_BAD_REQUEST)
        file_list = request.FILES.getlist(filekey)
        if len(file_list) > 1:
            return Response({"detail","Only one file in %s is acceptable." % (filekey)},status=status.HTTP_400_BAD_REQUEST)
        elif len(file_list) == 0:
            return Response({"detail","%s has no files." % (filekey)},status=status.HTTP_400_BAD_REQUEST)

        uploadedfile = file_list[0]
        filename = uploadedfile.name
        rootname,fileext = os.path.splitext(uploadedfile.name)

        if uploadedfile.multiple_chunks:
            for chunk in uploadedfile.chunks():
                    f.write(chunk)
        else:
            f.write(uploadedfile.read())

        return Response(data={'msg':'OK','URL':new_obj.url},status=status.HTTP_200_OK)

compound_ra_type_abbreviation = {'tc' : Compound.RAType.target,
                                 'sc': Compound.RAType.source}
compound_ra_type_code = {v: k for k, v in compound_ra_type_abbreviation.items()}



@method_decorator((csrf_protect,ensure_csrf_cookie), name='dispatch')
class CompoundView(GenericAPIView, CreateModelMixin, ListModelMixin):
    authentication_classes = [SessionAuthentication]
    permission_classes = [IsAuthenticated,IsProjectOwner]
    serializer_class = SlugCompoundCASRNSSerializer
    lookup_field = 'project'
    lookup_url_kwarg = 'project'
    
    def get_queryset(self):
        # The first thing that get(), post(), put() and delete() methods 
        # of a GenericView from Django REST framework is to call get_queryset()
        # directly or indirectly (the first thing that get_objects() method does
        # is calling get_queryset()). This way we are we are overwriting 
        # self.kwargs before they are used.
        if self.kwargs['ra_type'] in compound_ra_type_abbreviation:
            self.kwargs['ra_type'] = compound_ra_type_abbreviation[self.kwargs['ra_type']]
        elif self.kwargs['ra_type'] in compound_ra_type_abbreviation.values():
            pass
        else:
            return server_error(self.request)
        self.kwargs['project'] = int(self.kwargs['project'])
        return Compound.objects.filter(project=int(self.kwargs['project']),
                                        ra_type=self.kwargs['ra_type'])
    
    def get(self, request, project, ra_type):
        return self.list(request)
    def post(self, request, project, ra_type):
        db_ra_type = compound_ra_type_abbreviation[ra_type]
        data = dict(request.data)
        #for key in data:
        #    data[key] = data[key][0]


        with transaction.atomic():

            last_int_id = self.get_queryset().aggregate(Max('int_id'))['int_id__max']
            if (last_int_id is None):
                new_int_id = 1
            else:
                new_int_id = last_int_id + 1
            data['int_id'] = new_int_id
            data['ra_type'] = db_ra_type
            if ra_type == 'tc':
                data['tanimoto'] = 1.0
            serializer = CompoundSerializer(data=data)
            serializer.is_valid(raise_exception=True)
            try:
                self.perform_create(serializer)
            except IntegrityError as e:
                if str(e) == 'UNIQUE constraint failed: API_compound.project_id, API_compound.smiles':
                    message = 'Structure (SMILES) conflicts with an already saved one.'
                    raise serializers.ValidationError(message)
                else:
                    raise e
            if ra_type == 'tc':
                qcompounds = Compound.objects.filter(ra_type=db_ra_type,project=project,int_id=new_int_id)
                compound_id = qcompounds.values_list('id', flat=True)[0]
                tserializer= TCompoundSerializer(data={'project':project,'compound':compound_id})
                tserializer.is_valid(raise_exception=True)
                tserializer.save()
            else:
                #compute tanimoto here

                pass

            if data['cas_rn'] is None:
                data['cas_rn'] = []
            if len(data['cas_rn']) > 0:

                casrn_serializer = CompoundCASRNSerializer(many=True, data=[{"compound":compound_id, "cas_rn": casrn} for casrn in data['cas_rn']])
                casrn_serializer.is_valid(raise_exception=True)
                casrn_serializer.save()
            sserializer = SlugCompoundCASRNSSerializerNoValidation(data=data)
            sserializer.is_valid(raise_exception=True)
        headers = self.get_success_headers(sserializer.data)
        return Response(sserializer.data, status=status.HTTP_201_CREATED, headers=headers)

@method_decorator((csrf_protect,ensure_csrf_cookie), name='dispatch')
class CompoundCreateListView(GenericAPIView, CreateModelMixin, ListModelMixin):
    authentication_classes = [SessionAuthentication]
    permission_classes = [IsAuthenticated,IsProjectOwner]
    serializer_class = SlugCompoundCASRNSSerializer
    lookup_field = 'project'
    lookup_url_kwarg = 'project'
    

    def _perform_create(self, serializer, data, request, project):
            #compute tanimoto here (change serializer.validated_data and data)
            projects = set([compound['project'] for compound in serializer.validated_data])
            projects.add(project)
            tc_fingerprint = TCFingerPrint()
            TCs = Compound.objects.filter(project__in=projects,ra_type=compound_ra_type_abbreviation['tc'],int_id=1)
            TCs = TCs.values('project','smiles')
            for tc in TCs:
                p = tc['project']
                if tc_fingerprint.get(request, p) is None:
                    tc_fingerprint.set(request,p,tc['smiles'])
            for compound, compoundd in zip(serializer.validated_data, data):
                tanimoto = tc_fingerprint.similarity_from_smiles(request, compound['project'].id, compound['smiles'])
                compound['tanimoto'] = tanimoto
                compoundd['tanimoto'] = tanimoto
            for row in serializer.validated_data:
                row['project'] = row['project'].pk
            serializer = CompoundSerializer(data=serializer.validated_data, many=True)
            serializer.is_valid(raise_exception=True)
            try:
                self.perform_create(serializer)
            except IntegrityError as e:
                if str(e) == 'UNIQUE constraint failed: API_compound.project_id, API_compound.smiles':
                    smiles_list = [compound['smiles'] for compound in serializer.validated_data]
                    q_already_saves_compounds_smiles = Compound.objects.filter(project= self.kwargs['project'],smiles__in=smiles_list).values_list('smiles',flat=True)
                    duplicated_indexes = [i for i,smiles in enumerate(smiles_list) if smiles in q_already_saves_compounds_smiles]
                    del smiles_list
                    for index in duplicated_indexes:
                        serializer.validated_data.pop(index)
                        data.pop(index)
                    for row in serializer.validated_data:
                        row['project'] = row['project'].pk
                    serializer = CompoundSerializer(data=serializer.validated_data, many=True)
                    serializer.is_valid(raise_exception=True)
                    self._perform_create(serializer, data, request, project)
                elif str(e) == 'UNIQUE constraint failed: API_compound.project_id, API_compound.chembl_id':
                    chembl_id_list = [compound['chembl_id'] for compound in serializer.validated_data]
                    q_already_saves_compounds_chembl_id = Compound.objects.filter(project= self.kwargs['project'],chembl_id__in=chembl_id_list).values_list('chembl_id',flat=True)
                    duplicated_indexes = [i for i,chembl_id in enumerate(chembl_id_list) if chembl_id in q_already_saves_compounds_chembl_id]
                    del chembl_id_list
                    for index in duplicated_indexes:
                        serializer.validated_data.pop(index)
                        data.pop(index)
                    for row in serializer.validated_data:
                       row['project'] = row['project'].pk
                    serializer = CompoundSerializer(data=serializer.validated_data, many=True)
                    serializer.is_valid(raise_exception=True)
                    self._perform_create(serializer, data, request, project)
                else:
                    raise e
    def get_queryset(self):
        # The first thing that get(), post(), put() and delete() methods 
        # of a GenericView from Django REST framework is to call get_queryset()
        # directly or indirectly (the first thing that get_objects() method does
        # is calling get_queryset()). This way we are we are overwriting 
        # self.kwargs before they are used.
        if self.kwargs['ra_type'] in compound_ra_type_abbreviation:
            self.kwargs['ra_type'] = compound_ra_type_abbreviation[self.kwargs['ra_type']]
        elif self.kwargs['ra_type'] in compound_ra_type_abbreviation.values():
            pass
        else:
            return server_error(self.request)
        self.kwargs['project'] = int(self.kwargs['project'])
        return Compound.objects.filter(project=int(self.kwargs['project']),
                                        ra_type=self.kwargs['ra_type'])
    
    def get(self, request, project, ra_type):
        return self.list(request)
    def post(self, request, project, ra_type):
        #debug
        if ra_type == 'sc':
            Compound.objects.filter(ra_type=1,project=project).delete()
        
        db_ra_type = compound_ra_type_abbreviation[ra_type]

        with transaction.atomic():
            last_int_id = self.get_queryset().aggregate(Max('int_id'))['int_id__max']
            counter = 0
            data = list(request.data)
            int_id_list = []
            for row in data:
                if (last_int_id is None):
                    new_int_id = 1 + counter
                else:
                    new_int_id = last_int_id + counter + 1
                row['int_id'] = new_int_id
                int_id_list.append(new_int_id)
                row['ra_type'] = db_ra_type
                if row['cas_rn'] is None:
                    row['cas_rn'] = []
                counter += 1
            serializer = CompoundSerializer(data=data, many=True)
            serializer.is_valid(raise_exception=True)
        self._perform_create(serializer, data, request, project)
        with transaction.atomic():
            qcompounds = Compound.objects.filter(ra_type=db_ra_type,project=project,int_id__in=int_id_list)
            compound_ids = qcompounds.order_by('int_id').values_list('id', flat=True)
            tc_data = []
            casrn_data = []
            for compound_id, row in zip(compound_ids, data):
                casrn_data += [{"compound":compound_id, "cas_rn": casrn} for casrn in row['cas_rn']]
                if row['ra_type'] == compound_ra_type_abbreviation['tc']:
                    tc_data.append({'project':project,'compound':compound_id})
            if len(tc_data) > 0:
                tserializer = TCompoundSerializer(many=True, data=tc_data)
                tserializer.is_valid(raise_exception=True)
                tserializer.save()
            
            if len(casrn_data) > 0:
                casrn_serializer = CompoundCASRNSerializer(many=True, data=casrn_data)
                casrn_serializer.is_valid(raise_exception=True)
                casrn_serializer.save()
        sserializer = SlugCompoundCASRNSSerializerNoValidation(many=True, data=data)
        sserializer.is_valid(raise_exception=True)
        headers = self.get_success_headers(sserializer.data)
        return Response(sserializer.data, status=status.HTTP_201_CREATED, headers=headers)        

@method_decorator((csrf_protect,ensure_csrf_cookie), name='dispatch')    
class CompoundByIntIdView(RetrieveUpdateDestroyAPIView):
    authentication_classes = [SessionAuthentication]
    permission_classes = [IsAuthenticated,IsProjectOwner]
    serializer_class = SlugCompoundCASRNSSerializer
    lookup_field = 'int_id'
    lookup_url_kwarg = 'int_id'
    
    def get_queryset(self):
        # The first thing that get(), post(), put() and delete() methods 
        # of a GenericView from Django REST framework is to call get_queryset()
        # directly or indirectly (the first thing that get_objects() method does
        # is calling get_queryset()). This way we are we are overwriting 
        # self.kwargs before they are used.
        if self.kwargs['ra_type'] in compound_ra_type_abbreviation:
            self.kwargs['ra_type'] = compound_ra_type_abbreviation[self.kwargs['ra_type']]
        elif self.kwargs['ra_type'] in compound_ra_type_abbreviation.values():
            pass
        else:
            return server_error(self.request)
        self.kwargs['project'] = int(self.kwargs['project'])
        self.kwargs['int_id'] = int(self.kwargs['int_id'])
        
        return Compound.objects.filter(project=self.kwargs['project'],
                                     ra_type=self.kwargs['ra_type'])
        
    def update(self, request, *args, **kwargs):
        partial = kwargs.pop('partial', False)
        instance = self.get_object()
        data = dict(request.data)
        ra_type = self.kwargs['ra_type']
        data['ra_type'] = ra_type
        project = self.kwargs['project']
        data['project'] = project
        data['int_id'] = self.kwargs['int_id']
        if data['ra_type'] == compound_ra_type_abbreviation['tc']:
            data['tanimoto'] = 1.0
        else:
            tc_fingerprint = TCFingerPrint()
            tc = Compound.objects.get(project=project,ra_type=compound_ra_type_abbreviation['tc'],int_id=1)
            if tc_fingerprint.get(request, project) is None:
                tc_fingerprint.set(request,project,tc.smiles)
            data['tanimoto'] = tc_fingerprint.similarity_from_smiles(request, project, tc.smiles)
        compound_id = instance.pk
        with transaction.atomic():
            CompoundCASRN.objects.filter(compound=compound_id).delete()
            if data['cas_rn'] is None:
                data['cas_rn'] = []
            if len(data['cas_rn']) > 0:
                casrn_serializer = CompoundCASRNSerializer(many=True, data=[{"compound":compound_id, "cas_rn": casrn} for casrn in data['cas_rn']])
                casrn_serializer.is_valid(raise_exception=True)
                casrn_serializer.save()
            serializer = self.get_serializer(instance, data=data, partial=partial)
            serializer.is_valid(raise_exception=True)
            self.perform_update(serializer)
            old_ra_type = instance.ra_type
            new_ra_type = serializer.validated_data['ra_type']
            if old_ra_type != new_ra_type:
                if new_ra_type == compound_ra_type_abbreviation['tc']:
                    tserializer= TCompoundSerializer(data={'project':project,'compound':compound_id})
                    tserializer.is_valid(raise_exception=True)
                    tserializer.save()
                elif new_ra_type == compound_ra_type_abbreviation['sc']:
                    TCompound.objects.filter(compound=compound_id).delete()
            try:
                self.perform_update(serializer)
            except IntegrityError as e:
                if str(e) == 'UNIQUE constraint failed: API_compound.project_id, API_compound.smiles':
                    message = 'Structure (SMILES) conflicts with an already saved one.'
                    raise serializers.ValidationError(message)
                else:
                    raise e

        if getattr(instance, '_prefetched_objects_cache', None):
            # If 'prefetch_related' has been applied to a queryset, we need to
            # forcibly invalidate the prefetch cache on the instance.
            instance._prefetched_objects_cache = {}

        return Response(serializer.data)
    
    def delete(self, request, project, ra_type, int_id):
        with transaction.atomic():
            compounds = self.get_queryset().order_by('int_id')
            compounds_l = compounds.select_for_update(nowait=True)
            compounds_l = [q for q in compounds_l if q.int_id != self.kwargs['int_id']]
            # instance = self.get_object()
            # CompoundCASRN.objects.filter(compound=instance.id).delete()
            response = super().delete(request, project, ra_type, int_id)
            compounds_to_update = []
            for idx, q in enumerate(compounds_l):
                if (q.int_id != idx + 1):
                    q.int_id = idx + 1
                    compounds_to_update.append(q)
            Compound.objects.bulk_update(compounds_to_update,['int_id'])

            return response
chembl2data_matrix = {
    'assay_description': 'description',
    'standard_type': 'name',
    'standard_units': 'std_unit',
    'standard_value': 'std_value',
    'units': 'unit',
    'value': 'value',
    'assay_chembl_id': 'assay_id',
    'text_value': 'text_value',
    'assay_type': 'assay_type',
    
}        
@method_decorator((csrf_protect,ensure_csrf_cookie), name='dispatch')
class ChemblDataMatrixView(APIView):
    authentication_classes = [SessionAuthentication]
    permission_classes = [IsAuthenticated,IsProjectOwner]
    def get(self, request, project, ra_type, int_id, compound_init_id):
        # The first thing that get(), post(), put() and delete() methods 
        # of a GenericView from Django REST framework is to call get_queryset()
        # directly or indirectly (the first thing that get_objects() method does
        # is calling get_queryset()). This way we are we are overwriting 
        # self.kwargs before they are used.
        if self.kwargs['ra_type'] in compound_ra_type_abbreviation:
            self.kwargs['ra_type'] = compound_ra_type_abbreviation[self.kwargs['ra_type']]
        elif self.kwargs['ra_type'] in compound_ra_type_abbreviation.values():
            pass
        else:
            return server_error(self.request)
        respdata = {'msg': 'OK'}
        respdata['CSRF_TOKEN'] = get_token(request) 
        return JsonResponse(respdata, status=status.HTTP_200_OK)
    def post(self, request, project, ra_type, int_id=None):
        chembl2data_matrix = {
            'assay_description': 'description',
            'standard_type': 'name',
            'standard_units': 'std_unit',
            'standard_value': 'std_value',
            'units': 'unit',
            'value': 'value',
            'assay_chembl_id': 'assay_id',
            'text_value': 'text_value',
            'assay_type': 'assay_type',
            
        }
        
        chembl2data_matrix_assay_type = {
            'calculated_pc': DataMatrixFields.AssayType.calculated_pc,
            'A': DataMatrixFields.AssayType.bioactivity,
            'P': DataMatrixFields.AssayType.pc,
        }
        
        
        # The first thing that get(), post(), put() and delete() methods 
        # of a GenericView from Django REST framework is to call get_queryset()
        # directly or indirectly (the first thing that get_objects() method does
        # is calling get_queryset()). This way we are we are overwriting 
        # self.kwargs before they are used.
        if self.kwargs['ra_type'] in compound_ra_type_abbreviation:
            self.kwargs['ra_type'] = compound_ra_type_abbreviation[self.kwargs['ra_type']]
        elif self.kwargs['ra_type'] in compound_ra_type_abbreviation.values():
            pass
        else:
            return server_error(self.request)
        
        if 'chembl_id' in request.data and int_id is not None:
            if 'chembl_id' is not None:
                with transaction.atomic():
                    instance = Compound.objects.get(project=project,ra_type=self.kwargs['ra_type'],int_id=int(int_id))
                    serializer = CompoundSerializer(instance, data={'chembl_id':request.data['chembl_id']}, partial=True)
                    serializer.is_valid(raise_exception=True)
                    serializer.save()


        if 'data' in request.data:
            request_data = request.data['data']
        else:
            request_data = request.data
        if int_id is not None:
            data_list = [{'int_id':int_id,'data':request_data}]
        else:
            data_list = request_data
            
        i = 0
        with transaction.atomic():
            for rdata in data_list:
                current_int_id = rdata['int_id']
                if rdata['data'] is None:
                    continue
                serializer = ChemblDataMatrixSerializer(data=rdata['data'], many=True)
                try:
                    serializer.is_valid(raise_exception=True)
                except Exception as e:
                    raise e
                data = serializer.data
                units_set = set()
                standard_units = {}
                for col in data:
                    if col is None:
                        continue
                    unit = col['units']
                    std_unit = col['standard_units']
                    if std_unit is None:
                        std_unit = unit
                    if unit is None and std_unit is not None:
                        col['units'] = col['standard_units']
                        col['value'] = col['standard_value']
                    elif  unit is None and std_unit is None:
                        standard_units[None] = [None, None]
                    if unit not in standard_units:
                        if (col['value'] is not None and col['standard_value'] is not None) and col['standard_value'] != 0:
                            equivalence = col['value'] / col['standard_value']
                        else:
                            equivalence = None
                        standard_units[unit] = [std_unit, equivalence]
                    units_set.add(unit)
                    units_set.add(std_unit)
                    

                units = Unit.objects.filter(symbol__in=list(units_set))
                defined_units = set([unit.symbol for unit in units])
                
                undefined_units = units_set - defined_units
                undefined_std_units = set()
                std_units = set()
                for unit in undefined_units:
                    if unit in standard_units:
                        if standard_units[unit][0] is not None: 
                            if standard_units[unit][0] not in defined_units:
                                undefined_std_units.add(standard_units[unit][0])
                            std_units.add(standard_units[unit][0])
                        else:
                            std_units.add(None)
                            if None not in defined_units:
                                undefined_std_units.add(None)
                new_std_std_units_data = []
                new_std_units_data = []
                for unit in undefined_std_units:
                    name = unit
                    if unit is None:
                        name = 'dimensionless'
                    new_std_std_units_data.append({'name': name, 'std_unit':None, 'project': project})
                    new_std_units_data.append({'name': name, 'type':None, 'equivalence': 1.0, 'symbol': unit})
                unit_type_serializer = UnitTypeSerializer(data=new_std_std_units_data, many=True)
                unit_type_serializer.is_valid(raise_exception=True)
                unit_type_serializer.save()
                if None in undefined_std_units:
                    undefined_std_units_no_none = set(undefined_std_units)
                    undefined_std_units_no_none.remove(None)
                    undefined_std_units_no_none.add('dimensionless')
                else:
                    undefined_std_units_no_none = undefined_std_units
                
                new_std_std_units = UnitType.objects.filter(name__in=list(undefined_std_units_no_none))

                unit_type_std_unit_name2id = { unit.name: unit.id for unit in new_std_std_units }
                for unit in new_std_units_data:
                    unit['type'] = unit_type_std_unit_name2id[unit['name']]
                unit_std_serializer = UnitSerializer(data = new_std_units_data, many=True)
                unit_std_serializer.is_valid(raise_exception=True)
                unit_std_serializer.save()
                new_std_units = Unit.objects.filter(name__in=list(undefined_std_units_no_none))
                std_unit_name2id = { unit.name: unit.id for unit in new_std_units }
                partial_new_std_std_units_data = []
                for unit in new_std_units_data:
                    partial_new_std_std_units_data.append({'id': unit_type_std_unit_name2id[unit['name']],
                                                        'std_unit': std_unit_name2id[unit['name']],
                                                        'project': project})
                partial_unit_type_serializer = UnitTypeSerializer(data=partial_new_std_std_units_data, many=True, partial=True)
                partial_unit_type_serializer.is_valid(raise_exception=True)
                partial_unit_type_serializer.save()
                non_std_undefined_units = undefined_units - undefined_std_units
                saved_std_units = Unit.objects.filter(symbol__in=list(std_units))
                saved_std_units_symbol2id = {unit.symbol: unit.id for unit in saved_std_units}
                new_units_data = []
                for unit in non_std_undefined_units:
                    name = unit
                    if unit is None:
                        name = 'dimensionless'
                    new_units_data.append({'name': name, 'type':saved_std_units_symbol2id[standard_units[unit][0]],
                                            'equivalence': standard_units[unit][1], 'symbol': unit})
                
                unit_serializer = UnitSerializer(data = new_units_data, many=True)
                unit_serializer.is_valid(raise_exception=True)
                unit_serializer.save()
                saved_units = Unit.objects.filter(symbol__in=list(units_set))
                saved_units_symbol2id = {unit.symbol: unit.id for unit in saved_units}
                    
                compound_id = Compound.objects.get(project=project, ra_type=self.kwargs['ra_type'],
                                                    int_id=current_int_id).pk
                
                new_data_matrix_data = {'compound': compound_id, 'project': project}
                if not DataMatrix.objects.filter(**new_data_matrix_data).exists():
                    data_matrix_serializer = DataMatrixSerializer(data=new_data_matrix_data,many=False)
                    data_matrix_serializer.is_valid(raise_exception=True)
                    data_matrix_serializer.save()
                row_id = DataMatrix.objects.get(**new_data_matrix_data).pk
                data = data.copy()
                fields_to_keep = set()
                for field in chembl2data_matrix:
                    fields_to_keep.add(chembl2data_matrix[field])
                for col in data:
                    for old_field in chembl2data_matrix.keys():
                        new_field = chembl2data_matrix[old_field]
                        if new_field in {'unit','std_unit'}:
                            if col[old_field] is not None:
                                col[new_field] = saved_units_symbol2id[col[old_field]]
                            else:
                                col[new_field] = None
                            if old_field != new_field:
                                col.pop(old_field)
                        elif new_field in {'text_value'}:
                            if col[old_field] is None:
                                col[new_field] = col['activity_comment']
                                col.pop('activity_comment')
                            else:
                                col[new_field] = col[old_field]
                            if old_field != new_field:
                                col.pop(old_field)
                        elif old_field == 'assay_type':
                            chembl_assay_type = col[old_field]
                            col[new_field] = chembl2data_matrix_assay_type[chembl_assay_type]
                        elif new_field != old_field:
                            col[new_field] = col[old_field]
                            col.pop(old_field)
                    for field in list(col.keys()):
                        if field not in fields_to_keep:
                            col.pop(field)
                    col['row'] = row_id
                

                try:

                    data_matrix_fields_serializer = DataMatrixFieldsSerializer(data=data,many=True)
                    data_matrix_fields_serializer.is_valid(raise_exception=True)
                except Exception as e:
                    print(data)
                    raise e
                data_matrix_fields_serializer.save()
        if int_id is not None:
            return Response(data_matrix_fields_serializer.data, status=status.HTTP_201_CREATED)
        else:
            return Response({'msg':'OK'}, status=status.HTTP_201_CREATED)
        
@method_decorator((csrf_protect,ensure_csrf_cookie), name='dispatch')
class DataMatrixFieldsView(ListAPIView):
    authentication_classes = [SessionAuthentication]
    permission_classes = [IsAuthenticated,IsProjectOwner]
    serializer_class = DataMatrixFieldsReadSerializer
    lookup_field = 'int_id'
    lookup_url_kwarg = 'int_id'
    
    def get_queryset(self):
        # The first thing that get(), post(), put() and delete() methods 
        # of a GenericView from Django REST framework is to call get_queryset()
        # directly or indirectly (the first thing that get_objects() method does
        # is calling get_queryset()). This way we are we are overwriting 
        # self.kwargs before they are used.
        if self.kwargs['ra_type'] in compound_ra_type_abbreviation:
            self.kwargs['ra_type'] = compound_ra_type_abbreviation[self.kwargs['ra_type']]
        elif self.kwargs['ra_type'] in compound_ra_type_abbreviation.values():
            pass
        else:
            return server_error(self.request)
        self.kwargs['project'] = int(self.kwargs['project'])
        self.kwargs['int_id'] = int(self.kwargs['int_id'])
        
        q = DataMatrixFields.objects.annotate(int_id=F('row__compound__int_id'))
        q = q.annotate(ra_type=F('row__compound__ra_type'))
        q = q.annotate(project=F('row__project'))
        q = q.filter(project=self.kwargs['project'], ra_type=self.kwargs['ra_type'])
        q = q.select_related('unit').select_related('std_unit')
        return q
    
@method_decorator((csrf_protect,ensure_csrf_cookie), name='dispatch')
class DataMatrixView(ListAPIView):
    authentication_classes = [SessionAuthentication]
    permission_classes = [IsAuthenticated,IsProjectOwner]
    serializer_class = CompoundDataMatrixSerializer
    lookup_field = 'project'
    lookup_url_kwarg = 'project'
    queryset = Compound.objects.all()
    
    def get_queryset(self):
        # The first thing that get(), post(), put() and delete() methods 
        # of a GenericView from Django REST framework is to call get_queryset()
        # directly or indirectly (the first thing that get_objects() method does
        # is calling get_queryset()). This way we are we are overwriting 
        # self.kwargs before they are used.
        if self.kwargs['ra_type'] in compound_ra_type_abbreviation:
            self.kwargs['ra_type'] = compound_ra_type_abbreviation[self.kwargs['ra_type']]
        elif self.kwargs['ra_type'] in compound_ra_type_abbreviation.values():
            pass
        else:
            return server_error(self.request)
        self.kwargs['project'] = int(self.kwargs['project'])
        return super().get_queryset().filter(ra_type=self.kwargs['ra_type'],project=self.kwargs['project'])
    
@method_decorator((csrf_protect,ensure_csrf_cookie), name='dispatch')
class DataMatrixHeatmapView(GenericAPIView, ListModelMixin):
    authentication_classes = [SessionAuthentication]
    permission_classes = [IsAuthenticated,IsProjectOwner]
    serializer_class = CompoundDataMatrixSerializer
    lookup_field = 'project'
    lookup_url_kwarg = 'project'
    queryset = Compound.objects.all()

    def get_serializer_class(self):
        assert self.serializer_class is not None, (
            "'%s' should either include a `serializer_class` attribute, "
            "or override the `get_serializer_class()` method."
            % self.__class__.__name__
        )
        if 'assay_type' in self.kwargs:
            if self.kwargs['assay_type'] == 'similarity':
                return CompoundSerializer

        return self.serializer_class
    
    def get_queryset(self):
        # The first thing that get(), post(), put() and delete() methods 
        # of a GenericView from Django REST framework is to call get_queryset()
        # directly or indirectly (the first thing that get_objects() method does
        # is calling get_queryset()). This way we are we are overwriting 
        # self.kwargs before they are used.
        self.kwargs['project'] = int(self.kwargs['project'])
        return super().get_queryset().filter(project=self.kwargs['project']).order_by('-ra_type','-int_id')
    
    def _setup_y_values(self, compound, y_values, y_value_max_size):
        if compound['name'] is None:
            name = ''
        else:
            name = compound['name']
        comp = compound_ra_type_code[compound['ra_type']] + ':#' + str(compound['int_id'])+': '+name
        y_values.append(comp)
        comp_len = len(comp)
        if y_value_max_size < comp_len:
            y_value_max_size = comp_len
        return y_value_max_size


    def get(self,request, project, json, assay_type='bioactivity'):
        
        assay_types = {
            'bioactivity': {
                'value':[DataMatrixFields.AssayType.bioactivity],
                'title':" Min-max normalized activity",
            },
            'pc': {
                'value':[DataMatrixFields.AssayType.calculated_pc, DataMatrixFields.AssayType.pc],
                'title':" Min-max normalized Physicochemical property",
            },
            'similarity': {
                'title':"Tanimoto similarity (Morgan fingerprints)",
            }
        }
        
        # assay_types = {
        #     'bioactivity': {
        #         'value':[DataMatrixFields.AssayType.bioactivity],
        #         'title': None,
        #     },
        #     'pc': {
        #         'value':[DataMatrixFields.AssayType.calculated_pc, DataMatrixFields.AssayType.pc],
        #         'title': None,
        #     }
        # }
        
        chembl_molecular_species = {'ACID':1.0, 'NEUTRAL':0.45, 'ZWITTERION':0.65, 'BASE':0.0,None:np.nan,'N/A':np.nan}
        
        queryset = self.filter_queryset(self.get_queryset())

        # page = self.paginate_queryset(queryset)
        # if page is not None:
        #     serializer = self.get_serializer(page, many=True)
        #     return self.get_paginated_response(serializer.data)

        serializer = self.get_serializer(queryset, many=True)
        data = serializer.data
        
        del queryset
        del serializer
        
        x_values_set = set()
        y_values = []
        values_dict_list = []
        values_unit_dict_list = []
        description_dict_list = []
        name_dict_list = []
        x_value_2_assay_id_dict = {}
        if assay_type != 'similarity':
            colums_dict = {'x_value': [],'Compound': [],'Assay_ID': [], 'value': [], 'value_unit': [], 'description':[],'name':[], 'alpha2': []}
        else:
            colums_dict = {'x_value': [],'Compound': [], 'value': []}

        i = 0
        y_value_max_size = 0
        for compound in data:
            # if i > 30 and compound_ra_type_code[compound['ra_type']] == 'sc':
            #     continue
            values_dict = {}
            values_unit_dict = {}
            description_dict = {}
            name_dict = {}
            if assay_type != 'similarity':
                if len(compound['data_matrix']) > 0:
                    y_value_max_size = self._setup_y_values(compound, y_values, y_value_max_size)
                    if len(compound['data_matrix']) and assay_type != 'similarity':
                        for field in compound['data_matrix'][0]['data_matrix_fields']:
                            if field['assay_type'] not in assay_types[assay_type]['value']:
                                continue
                            if assay_type == 'pc':
                                x_value = field['name']+' ('+field['assay_id']+')'
                            else:
                                x_value = field['assay_id']
                            x_values_set.add(x_value)
                            x_value_2_assay_id_dict[x_value] = field['assay_id']
                            description_dict[field['assay_id']] = field['description']
                            name_dict[field['assay_id']] = field['name']
                            if (field['std_value'] is None):
                                if field['value'] is None:
                                    value = None
                                    unit = None
                                else:
                                    value = field['value']
                                    unit = field['unit']

                            else:
                                value = field['std_value']
                                unit = field['std_unit']
                            
                            if value is None:
                                if field['text_value'] is None:
                                    values_dict[field['assay_id']] = 'N/A'
                                    values_unit_dict[field['assay_id']] = 'N/A'
                                else:
                                    values_dict[field['assay_id']] = field['text_value']
                                    values_unit_dict[field['assay_id']] = field['text_value']
                            else:
                                if unit is None:
                                    unit_suffix = ''
                                else:
                                    unit_suffix = ' ' + unit
                                values_dict[field['assay_id']] = value
                                values_unit_dict[field['assay_id']] = str(value) + unit_suffix
                        values_dict_list.append(values_dict)
                        values_unit_dict_list.append(values_unit_dict)
                        description_dict_list.append(description_dict)
                        name_dict_list.append(name_dict)
            else:
                y_value_max_size = self._setup_y_values(compound, y_values, y_value_max_size)
                x_value = 'Similarity'
                x_values_set.add(x_value)
                values_dict_list.append({x_value : round(compound['tanimoto'],2)})
            i += 1
             
        del data
        x_values = sorted(list(x_values_set))
        compound_with_data_y_values = []
        if assay_type != 'similarity':
            for comp, values_dict, values_unit_dict, description_dict, name_dict in zip(y_values,
                    values_dict_list, values_unit_dict_list, description_dict_list, name_dict_list):
                colums_dict_comp = dict(colums_dict)
                assay_counter = len(set(values_dict.keys()).intersection(set([x_value_2_assay_id_dict[x_value] for x_value in x_values])))
                if assay_counter > 0:
                    for x_value in x_values:
                        assay_id = x_value_2_assay_id_dict[x_value]
                        if assay_id in values_dict:
                            z_value = values_dict[assay_id]
                            z_value_unit = values_unit_dict[assay_id]
                            description = escape(description_dict[assay_id])
                            description = escape(description_dict[assay_id])
                            name = name_dict[assay_id]
                            assay_counter += 1
                        else:
                            z_value = None
                        if z_value is None:
                            z_value = 'N/A'
                            z_value_unit = 'N/A'
                            description = 'N/A'
                            name = 'N/A'

                        if isinstance(z_value, str) and z_value != 'N/A' and assay_id != 'molecular_species':
                            alpha2 = 1.0
                        else:
                            alpha2 = 0
                        colums_dict['x_value'].append(x_value)
                        colums_dict['Compound'].append(comp)
                        colums_dict['Assay_ID'].append(assay_id)
                        colums_dict['value'].append(z_value)
                        colums_dict['value_unit'].append(z_value_unit)
                        colums_dict['description'].append(description)
                        colums_dict['name'].append(name)
                        colums_dict['alpha2'].append(alpha2)
                    compound_with_data_y_values.append(comp)
        else:
            for comp, values_dict, in zip(y_values, values_dict_list):
                x_value = list(values_dict.keys())[0]
                colums_dict['x_value'].append(x_value)
                colums_dict['Compound'].append(comp)
                colums_dict['value'].append(values_dict[x_value])
                compound_with_data_y_values.append(comp)
        del values_dict_list
        del values_unit_dict_list
        del description_dict_list
        del name_dict_list
                
        dataframe = pd.DataFrame(colums_dict)
        
        del colums_dict
        if assay_type != 'similarity':
            dataframe['fscaled_value'] = None
            for assay_id in [x_value_2_assay_id_dict[x_value] for x_value in x_values]:
                idxs = np.where(dataframe['Assay_ID'] == assay_id)[0]
                if dataframe.loc[idxs,'alpha2'].any() != 0 or assay_id == 'molecular_species':
                    dataframe.loc[idxs,'fscaled_value'] = 'N/A'
                    if assay_id == 'molecular_species':
                        array = dataframe.loc[idxs,'value']
                        for species, value in chembl_molecular_species.items():
                            if species is None:
                                array.replace(to_replace='None',value=value,regex=False,inplace=True)
                                array.fillna(value=value,inplace=True)
                            else:
                                array.replace(to_replace=species,value=value,regex=False,inplace=True)
                        dataframe.loc[idxs,'fscaled_value'] = array
                    continue
                array = dataframe.loc[idxs,'value'].replace('N/A', np.nan).to_numpy(dtype=np.float32)
                # array_min = np.nanmin(array)
                array_min = 0
                array_max = np.nanmax(np.abs(array))
                array_max_min = array_max - array_min
                if array_max_min == 0:
                    dataframe.loc[idxs[dataframe.loc[idxs,'value'] != 'N/A'],'fscaled_value'] = 0.0
                else:
                    array_fscaled = (array - array_min)*1/array_max_min
                    dataframe.loc[idxs,'fscaled_value'] = array_fscaled
            dataframe['fscaled_value'] = dataframe['fscaled_value'].replace(np.nan,'N/A')

        if len(x_values) < 1 or len(y_values) < 1:
            return Response({"item": None,'heatmap_div_id': None,'status': 'No data'})
        if assay_type == 'similarity':
            map_color_field = 'value'
        else:
            map_color_field = 'fscaled_value'
        #Map colors
        mapper = LinearColorMapper(palette=bokeh.palettes.RdYlBu[10], low = 0,
                         high = max([i for i in dataframe[map_color_field].to_list() if isinstance(i, numbers.Number) and not isinstance(i, bool)]))
    
        y_values = compound_with_data_y_values
        # Define a figure

        #pan=PanTool(dimensions="width")
        #mytools = ["hover","tap","save","reset","wheel_zoom",pan]

        mytools = ["hover","tap","save","reset","wheel_zoom","pan"]
        #w=int(len(df_t.columns)*40,)
        cw=275
        ch=30
        #h=int(((w-cw)*len(df_t.index)/len(df_t.columns))+ch)
        w= 1024
        h= 132
        if (len(y_values) < 6 or len(x_values) < 6):
            min_border_right=100
            p = figure(
            plot_width=w,
            min_border_right=min_border_right,
            #title="Example freq",
            y_range=y_values,
            x_range = x_values,
            tools=mytools, 
            x_axis_location="above",
            active_drag=None,
            toolbar_location="left",
            toolbar_sticky = False
            )
        elif assay_type == 'pc':
            width = 60*len(x_values)
            if width < w:
                width = w
            height = 80*len(y_values)
            min_border_right=150
            p = figure(
                plot_width=width,
                plot_height= height,
                min_border_right=min_border_right,
                y_range= y_values,
                x_range = x_values,
                tools=mytools, 
                x_axis_location="above",
                active_drag=None,
                toolbar_location="left",
                toolbar_sticky = False
                )
        else:
            width = 18*len(x_values)
            if width < w:
                width = w
            height = 80*len(y_values)
            min_border_right=100
            p = figure(
                plot_width=width,
                plot_height= height,
                min_border_right=min_border_right,
                #title="Example freq",
                y_range= y_values,
                x_range = x_values,
                tools=mytools, 
                x_axis_location="above",
                active_drag=None,
                toolbar_location="left",
                toolbar_sticky = False
                )

        # Create rectangle for heatmap
        mysource = ColumnDataSource(dataframe)
        del dataframe


        if assay_type == 'pc' or assay_type == 'similarity':
            dummyw = 1024
        else:
            dummyw = w

        dummysource = ColumnDataSource({'Column':[0],'Row':[0]})
        p.rect(
            y='Column', 
            x='Row', 
            width=dummyw, 
            height=1, 
            source=dummysource,
            line_color=None, 
            fill_color=None,

            # set visual properties for selected glyphs
            selection_line_color=None,
            selection_fill_color=None,
            # set visual properties for non-selected glyphs
            nonselection_fill_alpha=0,
            nonselection_fill_color=None,
            nonselection_line_color=None

        )


        p.rect(
            y='Compound', 
            x='x_value', 
            width=0.9, 
            height=1, 
            source=mysource,
            line_color=None, 
            fill_color=transform(map_color_field, mapper),

            # set visual properties for selected glyphs
            selection_line_color="crimson",
            selection_fill_color=transform(map_color_field, mapper),
            # set visual properties for non-selected glyphs
            nonselection_fill_alpha=1,
            nonselection_fill_color=transform(map_color_field, mapper),
            nonselection_line_color=None

            )
        if assay_type != 'similarity':
            p.rect(
                y='Compound', 
                x='x_value', 
                width=0.9, 
                height=1, 
                source=mysource,
                line_color=None, 
                fill_color='black',
                fill_alpha='alpha2',
                
                # set visual properties for selected glyphs
                selection_line_color="crimson",
                selection_fill_color='black',
                selection_fill_alpha='alpha2',
                # set visual properties for non-selected glyphs
                nonselection_fill_alpha='alpha2',
                nonselection_fill_color='black',
                nonselection_line_color=None

                )

        p.title.text = assay_types[assay_type]['title']
        p.title.align = "left"
        p.title.text_font_size = "25px"
        p.add_layout(Title(text="Compounds", align="right"), "left")
        if assay_type != 'similarity':
            p.add_layout(Title(text="Assays", align="left"), "above")

        # Add legend

        color_bar = ColorBar(color_mapper=mapper, ticker=BasicTicker(), 
                             formatter=PrintfTickFormatter(format=r'%.3f'),
                             orientation="vertical", label_standoff=12, location=(0,0),
                             major_label_text_font_size = "9pt",major_tick_line_color="black")
        
        
        p.add_layout(color_bar, 'left')


        p.axis.axis_line_color = None
        p.axis.major_tick_line_color = None
        p.xaxis.major_label_text_font_size = "9pt"
        p.yaxis.major_label_text_font_size = "8pt"
        p.axis.major_label_standoff = 0
        p.xaxis.major_label_orientation = 1#"vertical"
        
        TOOLTIPS = """
            <div style="max-width:20%">
            <div>
                <span href="#" data-toggle="tooltip" title="Compound">@Compound</span>
            </div>
            <div>
                <span href="#" data-toggle="tooltip" title="Assay ID">@Assay_ID</span>             
            </div>
            <div>
                <span href="#" data-toggle="tooltip" title="Value">@value_unit</span>
            </div><div>
              <span href="#" data-toggle="tooltip" title="Description">@description</span>
            </div><div>
                <span href="#" data-toggle="tooltip" title="Assay type">@name</span>                
            </div>
            </div>
        """
        
        if assay_type != 'similarity':
            #Hover tool:
            p.select_one(HoverTool).tooltips = [
            ('Compound', '@Compound'),
            ('Assay ID', '@Assay_ID'),
            ('Value', '@value_unit'),
            ('Description','@description{safe}'),
            ('Assay type', '@name')
            ]
        else:
            p.select_one(HoverTool).tooltips = [
            ('Compound', '@Compound'),
            ('Value', '@value{0.00}'),
            ]
            
        #p.select_one(HoverTool).tooltips = TOOLTIPS
    
        
        script, div = components(p)
        heatmap_div_id = "heatmap_datamatrix_"+assay_type+"_project_"+str(self.kwargs['project'])
        json_p = json_item(p,heatmap_div_id)
        json_p['doc']['title'] = heatmap_div_id
        if json is None:
            context={
                'script' : script , 
                'div' : div,
            }
            return render(request, 'API/datamatrix.html', context)
        else:
            root = ET.fromstring("<document>"+script+"</document>")
            scripts= []
            for tag in root.findall('script'):
                 scripts.append(tag.text)
            return Response({"item": json_p,'heatmap_div_id': heatmap_div_id,'status':'OK'})

def get_filename_from_url(url):
    return os.path.basename(urlparse(url).path)

def is_url(url):
    """
    Not to be used for actually validating a url, but in our use case we only 
    care if it's a url or a file path, and they're pretty distinguishable
    """
    parts = urlparse(url)
    return all([parts.scheme, parts.netloc, parts.path])

def fetch_image(url):
    """
    Attempts to fetch an image from a url. 
    If successful returns a bytes object, else returns None
    :return:
    """
    try:
        with urllib.request.urlopen(url) as response:
            # security flaw?
            return io.BytesIO(response.read())
    except urllib.error.URLError:
        return None


@method_decorator((csrf_protect,ensure_csrf_cookie), name='dispatch')
class GenerateReportDocx(APIView):
    authentication_classes = [SessionAuthentication]
    permission_classes = [IsAuthenticated,IsProjectOwner]
    class _CKEditorTableHTMLParser(HTMLParser):
        html = ''
        def __init__(self,*args,**kwargs):
            self.html = ''
            super().__init__(*args,convert_charrefs=False,**kwargs)
        def handle_starttag(self, tag, attrs):
            if tag != 'tbody':
                self.html += '<'+tag+' '
                for attr in attrs:
                    self.html += attr[0]+'="'+attr[1]+'" '
                self.html += '>'
        def handle_endtag(self, tag):
            if tag != 'tbody':
                self.html += '</'+tag+'>'

        def handle_data(self, data):
            self.html += data
        def get_parsed_html(self):
            return self.html


    class CustomHtmlToDocx(HtmlToDocx):
        datas = []
        skip_writting_data = False
        word_style = None
        heading_starting_level = 0

        def __init__(self,*args,styles=None, heading_starting_level=0,**kwargs):
            self.heading_starting_level = heading_starting_level
            self.styles = styles
            self.base54re = base54re
            super().__init__(*args,**kwargs)

        def _add_hyperlink(self, paragraph, url, text, color=None, underline=True):
            """
            A function that places a hyperlink within a paragraph object.

            :param paragraph: The paragraph we are adding the hyperlink to.
            :param url: A string containing the required url
            :param text: The text displayed for the url
            :return: The hyperlink object
            """

            # This gets access to the document.xml.rels file and gets a new relation id value
            part = paragraph.part
            r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

            # Create the w:hyperlink tag and add needed values
            hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
            hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

            # Create a w:r element
            new_run = docx.oxml.shared.OxmlElement('w:r')

            # Create a new w:rPr element
            rPr = docx.oxml.shared.OxmlElement('w:rPr')

            # Add color if it is given
            if not color is None:
                c = docx.oxml.shared.OxmlElement('w:color')
                c.set(docx.oxml.shared.qn('w:val'), color)
                rPr.append(c)

            # Remove underlining if it is requested
            if not underline:
                u = docx.oxml.shared.OxmlElement('w:u')
                u.set(docx.oxml.shared.qn('w:val'), 'none')
                rPr.append(u)

            # Join all the xml elements together add add the required text to the w:r element
            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)

            paragraph._p.append(hyperlink)
            return hyperlink
        def handle_img(self, current_attrs):
            if not self.include_images:
                self.skip = True
                self.skip_tag = 'img'
                return
            src = current_attrs['src']
            # fetch image
            src_is_url = is_url(src)
            if src_is_url:
                try:
                    image = fetch_image(src)
                except urllib.error.URLError:
                    image = None
            else:
                image = src
            # add image to doc
            if image:
                try:
                    if isinstance(self.doc, docx.document.Document):
                        self.doc.add_picture(image)
                    else:
                        self.add_image_to_cell(self.doc, image)
                except FileNotFoundError:
                    image = None
            if not image:
                image = src
                m = self.base54re.match(src)
                if m:
                    mimetype = m.group(1)
                    base64_bytes = m.group(2)
                    image_bytes = base64.b64decode(base64_bytes)
                    f = tempfile.NamedTemporaryFile(mode='wb', suffix=mimetypes.guess_extension(mimetype))
                    f.write(image_bytes)
                    f.flush()
                    try:
                        if isinstance(self.doc, docx.document.Document):
                            self.doc.add_picture(f.name)
                        else:
                            self.add_image_to_cell(self.doc, f.name)
                    except FileNotFoundError:
                        image = None
                    except Exception as e:
                        raise e
                    f.close()

                elif not image:
                    if src_is_url:
                        self.doc.add_paragraph("<image: %s>" % src)
                    else:
                        # avoid exposing filepaths in document
                        self.doc.add_paragraph("<image: %s>" % get_filename_from_url(src))

        def handle_starttag(self, tag, attrs):
            tag_none = False
            if self.skip:
                return
            if tag == 'a':
                self.skip_writting_data = True
            elif tag[0] == 'h' and len(tag) == 2:
                if isinstance(self.doc, docx.document.Document):
                    h_size = int(tag[1])+self.heading_starting_level
                    self.paragraph = self.doc.add_heading(level=min(h_size, 9))
                else:
                    self.paragraph = self.doc.add_paragraph()
                tag_none = True
            elif tag == 'figcaption':
                self.paragraph = self.doc.add_paragraph()
                self.word_style='Caption'
            elif tag == 'blockquote':
                self.word_style='Quote'
            if tag_none:
                super().handle_starttag('¨', attrs)
            else:
                super().handle_starttag(tag, attrs)

        def handle_endtag(self, tag):
            
            if self.skip:
                if not tag == self.skip_tag:
                    return
                
                if self.instances_to_skip > 0:
                    self.instances_to_skip -= 1
                    return
                self.skip = False
                self.skip_tag = None
                self.paragraph = None
            
            if tag == 'a':
                self.skip_writting_data = False
                link = self.tags.pop(tag)
                href = link['href']
                if len(self.datas) > 0:
                    text = self.datas.pop()
                else:
                    text = href
                
                self._add_hyperlink(self.paragraph,href,text)
                return
            elif tag == 'figcaption':
                self.word_style = None
            elif tag == 'blockquote':
                self.word_style = None


            if len(self.datas) > 0:
                self.datas.pop()
            super().handle_endtag(tag)
        def handle_data(self, data):
            if self.skip:
                return
            if not self.paragraph:
                self.paragraph = self.doc.add_paragraph()
            self.datas.append(data)
            if self.skip_writting_data:
                return
            if self.word_style is not None:
                self.paragraph.style = self.styles[self.word_style]
            super().handle_data(data)


    def _clean_ckeditor_html(self, html):
        parser = self._CKEditorTableHTMLParser()
        parser.feed(html)
        return parser.get_parsed_html()
    def _recursive_sections(self, document, sections_dict, data, heading=0):

        data_matrix_assay_type2readable = {
            DataMatrixFields.AssayType.calculated_pc: "Calculated PC",
            DataMatrixFields.AssayType.bioactivity: "Bioactivity",
            DataMatrixFields.AssayType.pc: "PC",
        }

        for section in sections_dict:
            if section['step'] is None:
                section_title = section['name']
            else:
                section_title = str(section['step'])+'.'+section['name']
            if section['type'] not in {"datatable","datamatrix_single"}:
                document.add_heading(section_title, heading)
            
            if len(section['subsections']) > 0:
                self._recursive_sections(document,section['subsections'],sections_dict, heading=heading+1)
            if section['description'] is not None:
                p = document.add_paragraph(section['description'])
            if section['type'] == "free-text":
                new_parser = self.CustomHtmlToDocx(heading_starting_level=heading-1,styles=document.styles)
                new_parser.add_html_to_document(self._clean_ckeditor_html(data[section['field']]), document)
            elif section['type'] == "compound":
                table = document.add_table(rows=1, cols=4)
                col0 = table.columns[0]
                col0.width = Cm(1.5)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '#'
                hdr_cells[1].text = 'CAS RN'
                hdr_cells[2].text = 'ChEMBL ID'
                hdr_cells[3].text = 'Structure'
                for compound in data[section['field']]:
                    srcs = glob(os.path.join(settings.MEDIA_ROOT_REPORTS,'images','compound_img_'+str(compound['id'])+'*'))
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(compound['int_id'])
                    row_cells[1].text = ', '.join(compound['cas_rn'])
                    row_cells[2].text = compound['chembl_id']
                    if len(srcs) > 0 :
                        paragraph = row_cells[3].paragraphs[0]
                        run = paragraph.add_run()
                        run.add_picture(srcs[0],height=1200000,width=1200000)
                    else:
                        row_cells[3].text = ''
            elif section['type'] == "datatable":
                document.add_section()
                docx_section = document.sections[-1]
                docx_section.orientation = WD_ORIENT.LANDSCAPE
                new_height = docx_section.page_width
                docx_section.page_width = docx_section.page_height
                docx_section.page_height = new_height
                document.add_heading(section_title, heading)
                nrows = len(data[section['field']]['Compound'])
                ncols = len(data[section['field']].keys())
                assay_headers = set(data[section['field']].keys())
                assay_headers.remove('Compound')
                headers = ['Compound'] + sorted(assay_headers)
                table = document.add_table(rows=1, cols=ncols)
                hdr_cells = table.rows[0].cells
                for hdr_cell, hdr_name in zip(hdr_cells,headers):
                    hdr_cell.text = hdr_name
                for row in range(0,nrows):
                    row_cells = table.add_row().cells
                    row_cells[0].text = data[section['field']]['Compound'][row]
                    for idx, hdr_name in enumerate(headers[1:]):
                        row_cells[idx+1].text = data[section['field']][hdr_name][row]
                document.add_section()
                docx_section2 = document.sections[-1]
                docx_section2.orientation = WD_ORIENT.PORTRAIT
                new_height = docx_section2.page_width
                docx_section2.page_width = docx_section2.page_height
                docx_section2.page_height = new_height
            elif section['type'] == "datamatrix_single":
                document.add_section()
                docx_section = document.sections[-1]
                docx_section.orientation = WD_ORIENT.LANDSCAPE
                new_height = docx_section.page_width
                docx_section.page_width = docx_section.page_height
                docx_section.page_height = new_height
                document.add_heading(section_title, heading-1)
                nrows = len(data[section['field']])
                headers = ['Property','Value','Units','Description','Assay type','Assay ID']
                fields = ['name', 'std_value', 'std_unit','description', 'assay_type','assay_id']

                ncols = len(headers)
                table = document.add_table(rows=1, cols=ncols)
                hdr_cells = table.rows[0].cells
                for hdr_cell, hdr_name in zip(hdr_cells,headers):
                    hdr_cell.text = hdr_name
                for row in range(0,nrows):
                    row_cells = table.add_row().cells
                    for idx, field in enumerate(fields):
                        value = data[section['field']][row][field]
                        if field == 'assay_type':
                            row_cells[idx].text = data_matrix_assay_type2readable[value]
                        elif field == 'assay_id' and value[0:6] == 'CHEMBL':
                            row_cells[idx].text = ''
                            paragraph = row_cells[idx].paragraphs[0]
                            self.CustomHtmlToDocx._add_hyperlink(None,paragraph,"https://www.ebi.ac.uk/chembl/assay_report_card/"+value+"/",value)
                        else:
                            row_cells[idx].text = str(value)
                document.add_section()
                docx_section2 = document.sections[-1]
                docx_section2.orientation = WD_ORIENT.PORTRAIT
                new_height = docx_section2.page_width
                docx_section2.page_width = docx_section2.page_height
                docx_section2.page_height = new_height

    def get(self, request, project):
        with open(settings.SECTIONS_FILE_PATH) as f:
            sections_dict = json.load(f)
        document = Document()
        style = document.styles['Quote']
        style.paragraph_format.left_indent = Pt(24)

        document.add_heading('Read-Across report', 0)

        data_matrix_serializer = CompoundDataMatrixSerializer(Compound.objects.filter(project=project), many=True)
        data_matrix_data = data_matrix_serializer.data

        assay_types = {
            'bioactivity': {
                'value':[DataMatrixFields.AssayType.bioactivity],
                'title':" Min-max normalized activity",
            },
            'pc': {
                'value':[DataMatrixFields.AssayType.calculated_pc],
                'title':" Min-max normalized Physicochemical property",
            }
        } #, DataMatrixFields.AssayType.pc

        step2node_seq = {1: 1, 2: 2, 3: 4, 4: 5, 5: 6}

        for section in sections_dict['sections']:
            if section['step'] is None:
                section_title = section['name']
            else:
                section_title = str(section['step'])+'.'+section['name']
                q_comments = NodesModel.objects.filter(project=project,node_seq=step2node_seq[section['step']]).values('outputs_comments')
                node_comments = q_comments[0]['outputs_comments']
            if section['name'] != "Appendix: TC Physicochemical, ADME and Toxicity data":
                document.add_heading(section_title, 1)
            if section['step'] == 1:
                q = ProblemDescription.objects.filter(project=project).values()
                data = {'node-comments': node_comments}
                if len(q) > 0:
                    data = q[0]
                    data['node-comments'] = node_comments
                else:
                    continue
            elif section['step'] == 2:
                data = {'node-comments': node_comments,'compounds':[]}
                q = Compound.objects.filter(project=project,ra_type=Compound.RAType.target)

                if len(q) > 0:
                    data['compounds'] = SlugCompoundCASRNSSerializer(q,many=True).data


                else:
                    continue
            elif section['step'] == 3:
                data = {'node-comments': node_comments}
                q = InitialRAxHypothesis.objects.filter(project=project).values()
                if len(q) > 0:
                    data = q[0]
                    data['node-comments'] = node_comments
                else:
                    continue
            elif section['step'] == 4:
                data = {'node-comments': node_comments,'compounds':[]}
                q = Compound.objects.filter(project=project,ra_type=Compound.RAType.source)
                if len(q) > 0:
                    data['compounds'] = SlugCompoundCASRNSSerializer(q,many=True).data
                    xvalues = set()
                    data['pc'] = {'Compound':[]}
                    i = 0
                    for compound in data_matrix_data:
                        # if i > 10 and compound_ra_type_code[compound['ra_type']] == 'sc':
                        #     continue
                        if len(compound['data_matrix']) > 0:
                            for field in compound['data_matrix'][0]['data_matrix_fields']:
                                if field['assay_type'] not in assay_types['pc']['value']:
                                    continue
                                if field['assay_id'] not in {'cx_most_apka','cx_logd','cx_logp','mw_freebase','molecular_species','psa','qed_weighted'}:
                                    continue
                                x_value = field['name']
                                if x_value not in xvalues:
                                    data['pc'][x_value] = []
                                    xvalues.add(x_value)
                            i +=1
                    i = 0
                    for compound in data_matrix_data:
                        # if i > 10 and compound_ra_type_code[compound['ra_type']] == 'sc':
                        #     continue
                        if len(compound['data_matrix']) > 0:
                            if compound['name'] is None:
                                name = ''
                            else:
                                name = compound['name']
                            comp = compound_ra_type_code[compound['ra_type']] + ':#' + str(compound['int_id'])+': '+name
                            data['pc']['Compound'].append(comp)
                            current_xvalues = set()
                            if len(compound['data_matrix']) > 0:
                                for field in compound['data_matrix'][0]['data_matrix_fields']:
                                    if field['assay_type'] not in assay_types['pc']['value']:
                                        continue
                                    if field['assay_id'] not in {'cx_most_apka','cx_logd','cx_logp','mw_freebase','molecular_species','psa','qed_weighted'}:
                                        continue

                                    x_value = field['name']

                                    if field['std_unit'] is None:
                                        unit = ''
                                    else:
                                        unit = ' '+field['std_unit']
                                    if field['std_value'] is not None:
                                        data['pc'][x_value].append(str(field['std_value'])+unit)
                                    else:
                                        data['pc'][x_value].append('–')
                                    current_xvalues.add(x_value)
                                for x_value in xvalues - current_xvalues:
                                    data['pc'][x_value].append('–')
                                i += 1

                else:
                    continue
            elif section['step'] is None:
                if section['name'] == "Appendix: TC Physicochemical, ADME and Toxicity data":
                    tc_data_matrix_serializer = CompoundDataMatrixSerializer(Compound.objects.filter(project=project,ra_type=Compound.RAType.target), many=True)
                    tc_data_matrix_data = tc_data_matrix_serializer.data
                    tc_data = tc_data_matrix_data[0]['data_matrix'][0]['data_matrix_fields']
                    tc_data.sort(key=lambda x: x['assay_id'])
                    tc_data.sort(key=lambda x: x['name'])
                    tc_data.sort(key=lambda x: x['assay_type'])
                    data = {'tc_data': tc_data}

            self._recursive_sections(document, section['subsections'], data, heading=2)


        # p = document.add_paragraph('A plain paragraph having some ')
        # p.add_run('bold').bold = True
        # p.add_run(' and some ')
        # p.add_run('italic.').italic = True

        # document.add_heading('Heading, level 1', level=1)
        # document.add_paragraph('Intense quote', style='Intense Quote')

        # document.add_paragraph(
        #     'first item in unordered list', style='List Bullet'
        # )
        # document.add_paragraph(
        #     'first item in ordered list', style='List Number'
        # )

        # document.add_picture(os.path.join(settings.MEDIA_ROOT_REPORTS,'monty-truth.png'), width=Inches(1.25))

        # records = (
        #     (3, '101', 'Spam'),
        #     (7, '422', 'Eggs'),
        #     (4, '631', 'Spam, spam, eggs, and spam')
        # )

        # table = document.add_table(rows=1, cols=3)
        # hdr_cells = table.rows[0].cells
        # hdr_cells[0].text = 'Qty'
        # hdr_cells[1].text = 'Id'
        # hdr_cells[2].text = 'Desc'
        # for qty, id, desc in records:
        #     row_cells = table.add_row().cells
        #     row_cells[0].text = str(qty)
        #     row_cells[1].text = id
        #     row_cells[2].text = desc

        # document.add_page_break()

        os.makedirs(settings.MEDIA_ROOT_REPORTS,mode=settings.DIRECTORY_DOWNLOAD_PERMISSIONS,exist_ok=True)
        filename = 'report_'+str(project)+'.docx'
        document.save(os.path.join(settings.MEDIA_ROOT_REPORTS,filename))
        return redirect(os.path.join(settings.MEDIA_URL_REPORTS,filename))

@method_decorator((csrf_protect,ensure_csrf_cookie), name='dispatch')
class SaveReportCompoundImage(APIView):
    authentication_classes = [SessionAuthentication]
    permission_classes = [IsAuthenticated,IsProjectOwner]
    def post(self,request,project):
        image_dir = os.path.join(settings.MEDIA_ROOT_REPORTS,'images')
        os.makedirs(image_dir,exist_ok=True)
        serializer = CompoundImageImageSerializer(data=request.data,many=False)
        serializer.is_valid(raise_exception=True)
        for compound, img in zip(request.data['compounds'],request.data['images']):
            if compound['project'] != project:
                continue
            compound_id = compound['id']
            m = base54re.match(img)
            if m:
                mimetype = m.group(1)
                base64_bytes = m.group(2)
                image_bytes = base64.b64decode(base64_bytes)
                with open(os.path.join(image_dir,'compound_img_'+str(compound_id)+mimetypes.guess_extension(mimetype)),'bw') as f:
                    f.write(image_bytes)
                    f.flush()
        return Response({'Result':'OK'})

@method_decorator((csrf_protect,ensure_csrf_cookie), name='dispatch')
class GenerateReportJson(APIView):
    authentication_classes = [SessionAuthentication]
    permission_classes = [IsAuthenticated,IsProjectOwner]
    
    def _recursive_sections(self, sections_dict, heading=0):

        data_matrix_assay_type2readable = {
            DataMatrixFields.AssayType.calculated_pc: "Calculated PC",
            DataMatrixFields.AssayType.bioactivity: "Bioactivity",
            DataMatrixFields.AssayType.pc: "PC",
        }

        for section in sections_dict:
            if section['step'] is None:
                section_title = section['name']
            else:
                section_title = str(section['step'])+'.'+section['name']
            # if section['type'] not in {"datatable","datamatrix_single"}:
            #     document.add_heading(section_title, heading)
            section['section_title'] = section_title
            if len(section['subsections']) > 0:
                self._recursive_sections(section['subsections'], heading=heading+1)
            # if section['description'] is not None:
            #     p = document.add_paragraph(section['description'])
            # if section['type'] == "free-text":
            #     new_parser = self.CustomHtmlToDocx(heading_starting_level=heading-1,styles=document.styles)
            #     new_parser.add_html_to_document(self._clean_ckeditor_html(data[section['field']]), document)
            # elif section['type'] == "compound":
            #     table = document.add_table(rows=1, cols=4)
            #     col0 = table.columns[0]
            #     col0.width = Cm(1.5)
            #     hdr_cells = table.rows[0].cells
            #     hdr_cells[0].text = '#'
            #     hdr_cells[1].text = 'CAS RN'
            #     hdr_cells[2].text = 'ChEMBL ID'
            #     hdr_cells[3].text = 'Structure'
            #     for compound in data[section['field']]:
            #         srcs = glob(os.path.join(settings.MEDIA_ROOT_REPORTS,'images','compound_img_'+str(compound['id'])+'*'))
            #         row_cells = table.add_row().cells
            #         row_cells[0].text = str(compound['int_id'])
            #         row_cells[1].text = ', '.join(compound['cas_rn'])
            #         row_cells[2].text = compound['chembl_id']
            #         if len(srcs) > 0 :
            #             paragraph = row_cells[3].paragraphs[0]
            #             run = paragraph.add_run()
            #             run.add_picture(srcs[0],height=1200000,width=1200000)
            #         else:
            #             row_cells[3].text = ''
            # elif section['type'] == "datatable":
            #     document.add_section()
            #     docx_section = document.sections[-1]
            #     docx_section.orientation = WD_ORIENT.LANDSCAPE
            #     new_height = docx_section.page_width
            #     docx_section.page_width = docx_section.page_height
            #     docx_section.page_height = new_height
            #     document.add_heading(section_title, heading)
            #     nrows = len(data['pc']['Compound'])
            #     ncols = len(data['pc'].keys())
            #     assay_headers = set(data['pc'].keys())
            #     assay_headers.remove('Compound')
            #     headers = ['Compound'] + sorted(assay_headers)
            #     table = document.add_table(rows=1, cols=ncols)
            #     hdr_cells = table.rows[0].cells
            #     for hdr_cell, hdr_name in zip(hdr_cells,headers):
            #         hdr_cell.text = hdr_name
            #     for row in range(0,nrows):
            #         row_cells = table.add_row().cells
            #         row_cells[0].text = data['pc']['Compound'][row]
            #         for idx, hdr_name in enumerate(headers[1:]):
            #             row_cells[idx+1].text = data['pc'][hdr_name][row]
            #     document.add_section()
            #     docx_section2 = document.sections[-1]
            #     docx_section2.orientation = WD_ORIENT.PORTRAIT
            #     new_height = docx_section2.page_width
            #     docx_section2.page_width = docx_section2.page_height
            #     docx_section2.page_height = new_height
            # elif section['type'] == "datamatrix_single":
            #     document.add_section()
            #     docx_section = document.sections[-1]
            #     docx_section.orientation = WD_ORIENT.LANDSCAPE
            #     new_height = docx_section.page_width
            #     docx_section.page_width = docx_section.page_height
            #     docx_section.page_height = new_height
            #     document.add_heading(section_title, heading-1)
            #     nrows = len(data[section['field']])
            #     headers = ['Property','Value','Units','Description','Assay type','Assay ID']
            #     fields = ['name', 'std_value', 'std_unit','description', 'assay_type','assay_id']

            #     ncols = len(headers)
            #     table = document.add_table(rows=1, cols=ncols)
            #     hdr_cells = table.rows[0].cells
            #     for hdr_cell, hdr_name in zip(hdr_cells,headers):
            #         hdr_cell.text = hdr_name
            #     for row in range(0,nrows):
            #         row_cells = table.add_row().cells
            #         for idx, field in enumerate(fields):
            #             value = data[section['field']][row][field]
            #             if field == 'assay_type':
            #                 row_cells[idx].text = data_matrix_assay_type2readable[value]
            #             elif field == 'assay_id' and value[0:6] == 'CHEMBL':
            #                 row_cells[idx].text = ''
            #                 paragraph = row_cells[idx].paragraphs[0]
            #                 self.CustomHtmlToDocx._add_hyperlink(None,paragraph,"https://www.ebi.ac.uk/chembl/assay_report_card/"+value+"/",value)
            #             else:
            #                 row_cells[idx].text = str(value)
            #     document.add_section()
            #     docx_section2 = document.sections[-1]
            #     docx_section2.orientation = WD_ORIENT.PORTRAIT
            #     new_height = docx_section2.page_width
            #     docx_section2.page_width = docx_section2.page_height
            #     docx_section2.page_height = new_height

    def get(self, request, project):
        with open(settings.SECTIONS_FILE_PATH) as f:
            sections_dict = json.load(f)
        output_dict = {'sections':[]}
        data_matrix_serializer = CompoundDataMatrixSerializer(Compound.objects.filter(project=project), many=True)
        data_matrix_data = data_matrix_serializer.data

        data_matrix_assay_type2readable = {
            DataMatrixFields.AssayType.calculated_pc: "Calculated PC",
            DataMatrixFields.AssayType.bioactivity: "Bioactivity",
            DataMatrixFields.AssayType.pc: "PC",
        }
        assay_types = {
            'bioactivity': {
                'value':[DataMatrixFields.AssayType.bioactivity],
                'title':" Min-max normalized activity",
            },
            'pc': {
                'value':[DataMatrixFields.AssayType.calculated_pc],
                'title':" Min-max normalized Physicochemical property",
            }
        } #, DataMatrixFields.AssayType.pc

        step2node_seq = {1: 1, 2: 2, 3: 4, 4: 5, 5: 6}

        for section in sections_dict['sections']:
            output_section = section.copy()
            if section['step'] is None:
                section_title = section['name']
            else:
                section_title = str(section['step'])+'.'+section['name']
                q_comments = NodesModel.objects.filter(project=project,node_seq=step2node_seq[section['step']]).values('outputs_comments')
                node_comments = q_comments[0]['outputs_comments']
            # if section['name'] != "Appendix: TC Physicochemical, ADME and Toxicity data":
            #     document.add_heading(section_title, 1)
            output_section['section_title'] = section_title
            if section['step'] == 1:
                q = ProblemDescription.objects.filter(project=project).values()
                data = {'node-comments': node_comments}
                if len(q) > 0:
                    data = q[0]
                    data['node-comments'] = node_comments
                else:
                    continue
            elif section['step'] == 2:
                data = {'node-comments': node_comments,'compounds':[]}
                q = Compound.objects.filter(project=project,ra_type=Compound.RAType.target)

                if len(q) > 0:
                    data['compounds'] = SlugCompoundCASRNSSerializer(q,many=True).data


                else:
                    continue
            elif section['step'] == 3:
                data = {'node-comments': node_comments}
                q = InitialRAxHypothesis.objects.filter(project=project).values()
                if len(q) > 0:
                    data = q[0]
                    data['node-comments'] = node_comments
                else:
                    continue
            elif section['step'] == 4:
                data = {'node-comments': node_comments,'compounds':[]}
                q = Compound.objects.filter(project=project,ra_type=Compound.RAType.source)
                if len(q) > 0:
                    data['compounds'] = SlugCompoundCASRNSSerializer(q,many=True).data
                    xvalues = set()
                    data['pc'] = {'Compound':[]}
                    i = 0
                    for compound in data_matrix_data:
                        # if i > 10 and compound_ra_type_code[compound['ra_type']] == 'sc':
                        #     continue
                        if len(compound['data_matrix']) > 0:
                            for field in compound['data_matrix'][0]['data_matrix_fields']:
                                if field['assay_type'] not in assay_types['pc']['value']:
                                    continue
                                if field['assay_id'] not in {'cx_most_apka','cx_logd','cx_logp','mw_freebase','molecular_species','psa','qed_weighted'}:
                                    continue
                                x_value = field['name']
                                if x_value not in xvalues:
                                    data['pc'][x_value] = []
                                    xvalues.add(x_value)
                            i +=1
                    i = 0
                    for compound in data_matrix_data:
                        # if i > 10 and compound_ra_type_code[compound['ra_type']] == 'sc':
                        #     continue
                        if len(compound['data_matrix']) > 0:
                            if compound['name'] is None:
                                name = ''
                            else:
                                name = compound['name']
                            comp = compound_ra_type_code[compound['ra_type']] + ':#' + str(compound['int_id'])+': '+name
                            data['pc']['Compound'].append(comp)
                            current_xvalues = set()
                            if len(compound['data_matrix']) > 0:
                                for field in compound['data_matrix'][0]['data_matrix_fields']:
                                    if field['assay_type'] not in assay_types['pc']['value']:
                                        continue
                                    if field['assay_id'] not in {'cx_most_apka','cx_logd','cx_logp','mw_freebase','molecular_species','psa','qed_weighted'}:
                                        continue

                                    x_value = field['name']

                                    if field['std_unit'] is None:
                                        unit = ''
                                    else:
                                        unit = ' '+field['std_unit']
                                    if field['std_value'] is not None:
                                        data['pc'][x_value].append(str(field['std_value'])+unit)
                                    else:
                                        data['pc'][x_value].append('–')
                                    current_xvalues.add(x_value)
                                for x_value in xvalues - current_xvalues:
                                    data['pc'][x_value].append('–')
                                i += 1

                else:
                    continue
            elif section['step'] is None:
                if section['name'] == "Appendix: TC Physicochemical, ADME and Toxicity data":
                    tc_data_matrix_serializer = CompoundDataMatrixSerializer(Compound.objects.filter(project=project,ra_type=Compound.RAType.target), many=True)
                    tc_data_matrix_data = tc_data_matrix_serializer.data
                    tc_data = None
                    if len(tc_data_matrix_data) > 0:
                        if 'data_matrix' in tc_data_matrix_data[0]:
                            if len(tc_data_matrix_data[0]['data_matrix']) > 0:
                                if 'data_matrix' in tc_data_matrix_data[0]['data_matrix'][0]:
                                    tc_data = tc_data_matrix_data[0]['data_matrix'][0]['data_matrix_fields']
                                    tc_data.sort(key=lambda x: x['assay_id'])
                                    tc_data.sort(key=lambda x: x['name'])
                                    tc_data.sort(key=lambda x: x['assay_type'])
                                    for row in tc_data:
                                        row['assay_type'] = data_matrix_assay_type2readable[row['assay_type']]
                                else:
                                    continue
                            else:
                                continue
                        else:
                            continue
                    else:
                        continue
                    data = {'tc_data': tc_data}
            output_section['data'] = data
            self._recursive_sections(output_section['subsections'], heading=2)
            output_dict['sections'].append(output_section)
            
        return Response(output_dict, status.HTTP_200_OK)


        # p = document.add_paragraph('A plain paragraph having some ')
        # p.add_run('bold').bold = True
        # p.add_run(' and some ')
        # p.add_run('italic.').italic = True

        # document.add_heading('Heading, level 1', level=1)
        # document.add_paragraph('Intense quote', style='Intense Quote')

        # document.add_paragraph(
        #     'first item in unordered list', style='List Bullet'
        # )
        # document.add_paragraph(
        #     'first item in ordered list', style='List Number'
        # )

        # document.add_picture(os.path.join(settings.MEDIA_ROOT_REPORTS,'monty-truth.png'), width=Inches(1.25))

        # records = (
        #     (3, '101', 'Spam'),
        #     (7, '422', 'Eggs'),
        #     (4, '631', 'Spam, spam, eggs, and spam')
        # )

        # table = document.add_table(rows=1, cols=3)
        # hdr_cells = table.rows[0].cells
        # hdr_cells[0].text = 'Qty'
        # hdr_cells[1].text = 'Id'
        # hdr_cells[2].text = 'Desc'
        # for qty, id, desc in records:
        #     row_cells = table.add_row().cells
        #     row_cells[0].text = str(qty)
        #     row_cells[1].text = id
        #     row_cells[2].text = desc

        # document.add_page_break()

        os.makedirs(settings.MEDIA_ROOT_REPORTS,mode=settings.DIRECTORY_DOWNLOAD_PERMISSIONS,exist_ok=True)
        filename = 'report_'+str(project)+'.docx'
        document.save(os.path.join(settings.MEDIA_ROOT_REPORTS,filename))
        return redirect(os.path.join(settings.MEDIA_URL_REPORTS,filename))



@method_decorator((csrf_protect,ensure_csrf_cookie), name='dispatch')
class Overview(ListAPIView):
    authentication_classes = [SessionAuthentication]
    permission_classes = [IsAuthenticated,IsProjectOwner]
    serializer_class = CompoundDataMatrixCountSerializer
    lookup_field = 'project'
    lookup_url_kwarg = 'project'
    queryset = Compound.objects.all()

    def get_queryset(self):
        q = self.queryset.filter(**{self.lookup_field:self.kwargs[self.lookup_url_kwarg]})
        return q.order_by('ra_type').order_by('int_id')
